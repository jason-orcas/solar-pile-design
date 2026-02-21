"""Convert SPORK_User_Manual.md to a formatted DOCX document."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Logo path relative to this script
LOGO_PATH = Path(__file__).parent.parent / "streamlit_app" / "assets" / "bowman_logo.png"


def set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg_color="1F4E79", font_color=RGBColor(255, 255, 255)):
    """Style a table header row with background color and white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = font_color
                run.font.size = Pt(9)


def add_formatted_paragraph(doc, text: str, style="Normal"):
    """Add a paragraph with inline formatting (bold, italic, code)."""
    p = doc.add_paragraph(style=style)

    # Split on bold (**...**), italic (*...*), and inline code (`...`)
    # Process bold first, then italic, then code within each segment
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            # Check for inline code within bold
            code_parts = re.split(r'(`[^`]+`)', inner)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = p.add_run(cp[1:-1])
                    run.bold = True
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                else:
                    run = p.add_run(cp)
                    run.bold = True
        else:
            # Process italic within non-bold segments
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                    run = p.add_run(ip[1:-1])
                    run.italic = True
                else:
                    # Process inline code
                    code_parts = re.split(r'(`[^`]+`)', ip)
                    for cp in code_parts:
                        if cp.startswith('`') and cp.endswith('`'):
                            run = p.add_run(cp[1:-1])
                            run.font.name = "Consolas"
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
                        else:
                            run = p.add_run(cp)
    return p


def parse_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into a 2D list of cell values."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator rows (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows: list[list[str]]):
    """Add a formatted table to the document."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                # Handle bold in cell text
                parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(9)
                    else:
                        # Handle inline code
                        code_parts = re.split(r'(`[^`]+`)', part)
                        for cp in code_parts:
                            if cp.startswith('`') and cp.endswith('`'):
                                run = p.add_run(cp[1:-1])
                                run.font.name = "Consolas"
                                run.font.size = Pt(8)
                            else:
                                run = p.add_run(cp)
                                run.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

    # Style header row
    if len(rows) > 0:
        style_header_row(table.rows[0])

    # Set column widths proportionally
    doc.add_paragraph()  # spacing after table


def add_toc_field(doc):
    """Insert a Word auto-updating Table of Contents field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run._r.append(fld_char_begin)

    run2 = p.add_run()
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u '
        '</w:instrText>'
    )
    run2._r.append(instr)

    run3 = p.add_run()
    fld_char_separate = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
    )
    run3._r.append(fld_char_separate)

    run4 = p.add_run("Right-click and select 'Update Field' to populate this table of contents.")
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.italic = True

    run5 = p.add_run()
    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run5._r.append(fld_char_end)


def restart_list_numbering(doc, paragraph):
    """Force a numbered list paragraph to restart numbering at 1.

    Creates a new numbering instance in the document's numbering definitions
    that references the same abstract numbering as 'List Number' but overrides
    the start value to 1.
    """
    numbering_part = doc.part.numbering_part
    ct_numbering = numbering_part._element

    # Find the numId currently assigned to this paragraph
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return
    num_id_elem = numPr.find(qn('w:numId'))
    if num_id_elem is None:
        return
    current_num_id = num_id_elem.get(qn('w:val'))

    # Look up the abstractNumId for that numId
    abstract_num_id = '0'
    for num_elem in ct_numbering.findall(qn('w:num')):
        if num_elem.get(qn('w:numId')) == current_num_id:
            abs_ref = num_elem.find(qn('w:abstractNumId'))
            if abs_ref is not None:
                abstract_num_id = abs_ref.get(qn('w:val'))
            break

    # Determine next available numId
    all_num_ids = [
        int(n.get(qn('w:numId'), '0'))
        for n in ct_numbering.findall(qn('w:num'))
    ]
    next_num_id = max(all_num_ids, default=0) + 1

    # Create new <w:num> that overrides level 0 start to 1
    new_num = parse_xml(
        f'<w:num {nsdecls("w")} w:numId="{next_num_id}">'
        f'  <w:abstractNumId w:val="{abstract_num_id}"/>'
        f'  <w:lvlOverride w:ilvl="0">'
        f'    <w:startOverride w:val="1"/>'
        f'  </w:lvlOverride>'
        f'</w:num>'
    )
    ct_numbering.append(new_num)

    # Point this paragraph to the new numbering instance
    num_id_elem.set(qn('w:val'), str(next_num_id))


def add_header_footer(section, version_text: str, logo_path: Path):
    """Add Bowman logo header and footer with title + page number to a section.

    Matches the manually edited v1 format:
      Header: Bowman logo, top-left (~1.1" wide)
      Footer: "Bowman -- Solar Pile Optimization & Report Kit Product Manual vX.X"
              followed by bold page number, right-aligned
    """
    # --- Header with logo ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    if logo_path.exists():
        run = hp.add_run()
        run.add_picture(str(logo_path), width=Inches(1.12))
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.space_before = Pt(0)

    # --- Footer with title text + page number ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""

    # Extract just the version number (e.g., "1.1" from "Version 1.1  |  February 2026")
    ver_match = re.search(r'(\d+\.\d+)', version_text)
    ver_num = ver_match.group(1) if ver_match else "1.1"

    # Footer text: "Bowman -- Solar Pile Optimization & Report Kit Product Manual vX.X"
    footer_text = f"Bowman \u2014 Solar Pile Optimization & Report Kit Product Manual v{ver_num}"
    run_text = fp.add_run(footer_text)
    run_text.font.name = "Calibri"
    run_text.font.size = Pt(10)

    # Add spaces between text and page number
    spacer = fp.add_run("  " * 20)
    spacer.font.size = Pt(10)

    # Page number field (bold)
    run_page = fp.add_run()
    fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_page._r.append(fld_begin)

    run_instr = fp.add_run()
    instr_text = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    )
    run_instr._r.append(instr_text)

    run_sep = fp.add_run()
    fld_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_sep._r.append(fld_sep)

    run_num = fp.add_run("1")
    run_num.bold = True
    run_num.font.size = Pt(10)

    run_end = fp.add_run()
    fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end._r.append(fld_end)

    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.space_before = Pt(0)


def convert_md_to_docx(md_path: str, docx_path: str):
    """Convert the SPORK User Manual markdown to a formatted DOCX."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    lines = md_text.split('\n')

    # Extract version text early for header/footer
    version_text = "Version 1.1  |  February 2026"
    for mline in lines[:10]:
        if mline.strip().startswith("Version"):
            version_text = mline.strip()
            break

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # --- Header and footer (logo + title + page number) ---
    add_header_footer(section, version_text, LOGO_PATH)

    # --- Define styles ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        if level == 1:
            heading_style.font.size = Pt(22)
            heading_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 3:
            heading_style.font.size = Pt(13)
            heading_style.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        elif level == 4:
            heading_style.font.size = Pt(11)
            heading_style.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    # --- Title page ---
    # Reduced from 6 blank lines to 4 since header logo now occupies top space
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("SPORK")
    run.font.size = Pt(42)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Solar Pile Optimization & Report Kit")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()

    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = type_p.add_run("User Manual")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_p.add_run(version_text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page break after title
    doc.add_page_break()

    # Track whether previous line was a numbered list item (for restart detection)
    prev_was_numbered = False

    # --- Parse and render content ---
    i = 0
    # Skip the markdown title block (first few lines)
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('## Table of Contents'):
            # Insert Word auto-TOC field referencing Heading styles
            doc.add_heading('Table of Contents', level=1)
            add_toc_field(doc)
            # Skip past the markdown TOC entries
            i += 1
            while i < len(lines) and lines[i].strip():
                i += 1
            doc.add_page_break()
            prev_was_numbered = False
            continue

        # Skip the initial title lines (already rendered on title page)
        if line.startswith('# SPORK User Manual') or line.strip() in (
            '**Solar Pile Optimization & Report Kit**',
        ) or line.strip().startswith('Version '):
            i += 1
            continue

        # Horizontal rules -> page break before major sections
        if line.strip() == '---':
            prev_was_numbered = False
            i += 1
            continue

        # Headings
        if line.startswith('####'):
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=4)
            prev_was_numbered = False
            i += 1
            continue
        elif line.startswith('###'):
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=3)
            prev_was_numbered = False
            i += 1
            continue
        elif line.startswith('## '):
            text = line.lstrip('#').strip()
            # Major sections get a page break
            if re.match(r'^\d+\.', text):
                doc.add_page_break()
            doc.add_heading(text, level=2)
            prev_was_numbered = False
            i += 1
            continue

        # Italic caption lines (e.g., *Page 02: Soil Profile*)
        if line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.size = Pt(10)
            prev_was_numbered = False
            i += 1
            continue

        # Tables
        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            add_table_to_doc(doc, rows)
            prev_was_numbered = False
            continue

        # Code blocks (indented with spaces)
        if line.strip().startswith('```'):
            i += 1  # skip opening ```
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            # Add as formatted code
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                run = p.add_run(cl)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
            prev_was_numbered = False
            continue

        # Indented code/formula lines (4+ spaces)
        if line.startswith('    ') and line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line.strip())
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
            prev_was_numbered = False
            i += 1
            continue

        # Bullet/list items
        if line.strip().startswith('- '):
            text = line.strip()[2:]
            add_formatted_paragraph(doc, text, style='List Bullet')
            prev_was_numbered = False
            i += 1
            continue

        # Numbered list items
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = add_formatted_paragraph(doc, text, style='List Number')
            if not prev_was_numbered:
                restart_list_numbering(doc, p)
            prev_was_numbered = True
            i += 1
            continue

        # Block quotes (> text)
        if line.strip().startswith('>'):
            text = line.strip().lstrip('>').strip()
            p = add_formatted_paragraph(doc, text)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add a left border effect via italic
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            prev_was_numbered = False
            i += 1
            continue

        # Regular paragraphs
        text = line.strip()
        if text:
            add_formatted_paragraph(doc, text)
            prev_was_numbered = False
        else:
            # Blank lines also break numbered list continuity
            prev_was_numbered = False
        i += 1

    # --- Save ---
    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    base = Path(__file__).parent
    convert_md_to_docx(
        str(base / "SPORK_User_Manual.md"),
        str(base / "SPORK_User_Manual_v1.2.docx"),
    )
